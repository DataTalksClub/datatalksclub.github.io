#!/usr/bin/env python
# coding: utf-8

import re

from io import BytesIO
from urllib import request

import docx


def clean_line(line):
    line = line.strip()
    line = line.strip('\uFEFF')
    return line


ts_two_digit_pattern = re.compile(r'^\[?(\d+):(\d+)\]?$')
ts_three_digit_pattern = re.compile(r'^\[?(\d+):(\d+):(\d+)\]?$')


def try_parse_time(line):
    match = ts_two_digit_pattern.match(line)

    if match is not None:
        m, s = int(match.group(1)), int(match.group(2))
        h = m // 60
        m = m % 60
        return h, m, s

    match = ts_three_digit_pattern.match(line)
    if match is not None:
        h = int(match.group(1))
        m = int(match.group(2))
        s = int(match.group(3))
        return h, m, s

    return None


def format_time(h, m, s):
    if h > 0:
        return '%d:%02d:%02d' % (h, m, s)
    return '%d:%02d' % (m, s)


def url_as_bytes_stream(url):
    with request.urlopen(url) as resp:
        buffer = resp.read()
    stream = BytesIO(buffer)
    return stream


def extract_lines(paragraphs):
    styled_lines = []

    for p in paragraphs:
        style = p.style.name.lower()
        paragraph = clean_line(p.text)

        if len(paragraph) == 0:
            continue

        if style != 'normal':
            paragraph = re.sub(r'\s+', ' ', paragraph)
            styled_lines.append((paragraph, style))
            continue

        for line in paragraph.split('\n'):
            line = clean_line(line)
            if len(line) == 0:
                continue
            styled_lines.append((line, style))

    return styled_lines


def read_docx(input_file):
    if input_file.startswith('http'):
        stream = url_as_bytes_stream(input_file)
        doc = docx.Document(stream)
    else:
        doc = docx.Document(input_file)

    h = None
    m = None
    s = None
    total_sec = None
    speaker = None
    after_time = False

    elements = []

    styled_lines = extract_lines(doc.paragraphs)

    for line, style in styled_lines:
        if style.startswith('heading'):
            element = {
                'header': line
            }
            elements.append(element)
            continue

        assert style == 'normal'

        if after_time == True:
            after_time = False
            speaker = line
            continue

        maybe_time = try_parse_time(line)

        if maybe_time is not None:
            h, m, s = maybe_time
            total_sec = h * 60 * 60 + m * 60 + s
            after_time = True
            continue

        assert speaker is not None
        assert total_sec is not None

        element = {
            'time': format_time(h, m, s),
            'sec': total_sec,
            'who': speaker,
            'line': line
        }

        elements.append(element)

    return elements




